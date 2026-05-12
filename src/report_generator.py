"""Word report generation from parsed ALCO data."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt


class ReportGenerator:
    """Generate board-level Word reports from parsed ALCO data."""

    MAX_ACTION_LENGTH = 220
    REPORT_PERIOD_LABEL = "For the Period Ended"
    MAX_ANALYSIS_METRIC_ROWS = 8
    MAX_BOARD_METRIC_ROWS = 6
    MAX_TABLE_COLUMNS = 5
    MAX_TABLE_ROWS = 6

    def __init__(self, parsed_data: dict[str, Any], output_dir: str | Path = "."):
        self.data = parsed_data
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_board_report(self) -> Path:
        """Generate the board report output document."""
        doc = self._base_document()
        metadata = self.data["metadata"]

        doc.add_paragraph(metadata["bank_name"].upper())
        doc.add_paragraph("Board of Directors — Financial Summary Report")
        doc.add_paragraph(f"{self.REPORT_PERIOD_LABEL} {metadata['report_date']}")

        doc.add_heading("AT A GLANCE", level=1)
        highlights = self.data.get("document_highlights", [])[:5]
        self._add_simple_table(doc, ["Key Highlight"], [[item] for item in highlights] or [["No highlights extracted."]])

        self._add_section_block(doc, "REVENUE & EARNINGS", ["net_interest_margin", "loan_portfolio", "investment_portfolio"])
        self._add_section_block(doc, "LIQUIDITY & FUNDING", ["liquidity", "deposit_composition"])
        self._add_section_block(doc, "BALANCE SHEET & CAPITAL", ["balance_sheet", "key_metrics_ratios", "regulatory_summaries"])
        self._add_section_block(doc, "INTEREST RATE RISK", ["interest_rate_risk", "rate_sensitivity_gap"])
        self._add_section_block(doc, "ECONOMIC OUTLOOK", ["economic_outlook"])

        doc.add_heading("BOARD CONSIDERATIONS", level=1)
        for action in self._board_actions():
            doc.add_paragraph(action, style="List Bullet")

        output_path = self.output_dir / self._filename("Board_Report")
        doc.save(output_path)
        return output_path

    def generate_alco_analysis(self) -> Path:
        """Generate the ALCO board analysis output document."""
        doc = self._base_document()
        metadata = self.data["metadata"]

        doc.add_paragraph("BOARD OF DIRECTORS")
        doc.add_paragraph(metadata["bank_name"])
        doc.add_paragraph("ALCO Report Analysis")
        doc.add_paragraph(f"{self.REPORT_PERIOD_LABEL} {metadata['report_date']}")

        sections = [
            ("I.  Executive Summary", ["key_metrics_ratios", "regulatory_summaries"]),
            ("II.  Financial Performance", ["net_interest_margin", "loan_portfolio"]),
            ("III.  Capital", ["balance_sheet", "regulatory_summaries"]),
            ("IV.  Liquidity", ["liquidity", "deposit_composition"]),
            ("V.  Interest Rate Risk", ["interest_rate_risk", "rate_sensitivity_gap"]),
            ("VI.  Investment Portfolio", ["investment_portfolio"]),
            ("VII.  Loan Portfolio", ["loan_portfolio"]),
            ("VIII.  Peer / Comparative Indicators", ["key_metrics_ratios"]),
            ("IX.  Economic Environment", ["economic_outlook"]),
            ("X.  Board Action Items", []),
        ]

        for title, keys in sections:
            doc.add_heading(title, level=1)
            if keys:
                narrative = self._compose_section_summary(keys)
                doc.add_paragraph(narrative)
                rows = self._collect_metric_rows(keys)
                if rows:
                    self._add_simple_table(doc, ["Metric", "Observed Value"], rows[: self.MAX_ANALYSIS_METRIC_ROWS])

            if "Interest Rate Risk" in title:
                doc.add_heading("12-Month NIM Sensitivity", level=2)
                doc.add_paragraph(self._compose_section_summary(["rate_sensitivity_gap"]))
                doc.add_heading("Economic Value of Equity (EVE)", level=2)
                doc.add_paragraph(self._compose_section_summary(["interest_rate_risk"]))

            if "Board Action Items" in title:
                for item in self._board_actions():
                    doc.add_paragraph(item, style="List Bullet")

        output_path = self.output_dir / self._filename("ALCO_Board_Analysis")
        doc.save(output_path)
        return output_path

    @staticmethod
    def _base_document() -> Document:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        return doc

    def _filename(self, report_kind: str) -> str:
        meta = self.data["metadata"]
        return f"{meta['bank_code']}_{report_kind}_{meta['report_month_year']}.docx"

    def _add_section_block(self, doc: Document, heading: str, section_keys: list[str]) -> None:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(self._compose_section_summary(section_keys))

        rows = self._collect_metric_rows(section_keys)
        if rows:
            self._add_simple_table(doc, ["Metric", "Observed Value"], rows[: self.MAX_BOARD_METRIC_ROWS])
            return

        table = self._first_table_for_sections(section_keys)
        if table:
            headers = table["headers"]
            rows = table["rows"]
            if headers and rows:
                self._add_simple_table(
                    doc,
                    headers[: self.MAX_TABLE_COLUMNS],
                    [row[: self.MAX_TABLE_COLUMNS] for row in rows[: self.MAX_TABLE_ROWS]],
                )

    def _compose_section_summary(self, section_keys: list[str]) -> str:
        parts: list[str] = []
        for key in section_keys:
            section = self.data.get("sections", {}).get(key, {})
            title = section.get("title")
            points = section.get("key_points", [])
            pages = section.get("pages", [])
            if points:
                parts.append(f"{title}: {points[0]}")
            elif pages:
                parts.append(f"{title}: information captured from pages {', '.join(map(str, pages[:8]))}.")

        if not parts:
            return "No directly extractable content was identified for this section in the provided PDF."
        return " ".join(parts)

    def _collect_metric_rows(self, section_keys: list[str]) -> list[list[str]]:
        rows: list[list[str]] = []
        for key in section_keys:
            section = self.data.get("sections", {}).get(key, {})
            for metric in section.get("metrics", [])[:5]:
                rows.append([metric.get("metric", ""), metric.get("value", "")])
        return rows

    def _first_table_for_sections(self, section_keys: list[str]) -> dict[str, Any] | None:
        for key in section_keys:
            section = self.data.get("sections", {}).get(key, {})
            tables = section.get("tables", [])
            if tables:
                return tables[0]
        return None

    @staticmethod
    def _add_simple_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        for row_values in rows:
            row = table.add_row().cells
            padded = row_values + [""] * (len(headers) - len(row_values))
            for idx, value in enumerate(padded[: len(headers)]):
                row[idx].text = str(value)

    def _board_actions(self) -> list[str]:
        actions: list[str] = []
        seen: set[str] = set()
        sections = self.data.get("sections", {})
        for section in sections.values():
            for point in section.get("key_points", [])[:4]:
                low = point.lower()
                if self._is_actionable_point(low) and any(
                    flag in low for flag in ("below", "decline", "loss", "risk", "breach")
                ):
                    cleaned = self._normalize_action(point)
                    if cleaned not in seen:
                        actions.append(cleaned)
                        seen.add(cleaned)
                if len(actions) >= 6:
                    return actions

        if not actions:
            actions.append("Review ALCO policy compliance and concentration trends from this period's report.")
            actions.append("Confirm liquidity and capital buffers remain aligned with board-approved targets.")
        return actions

    @staticmethod
    def _normalize_action(text: str) -> str:
        cleaned = " ".join(text.split())
        return cleaned[: ReportGenerator.MAX_ACTION_LENGTH]

    @staticmethod
    def _is_actionable_point(text: str) -> bool:
        if any(token in text for token in ("sector #", "parcoupon", "walam", "history date", "account name")):
            return False
        return len(text.split()) >= 8
